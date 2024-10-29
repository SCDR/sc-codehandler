import docx
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement

import os
import json

from file_handler import FileHandler


# 处理Docx文件
class DocxHandler:

    def __init__(
        self,
        name=None,
        pageSize=None,
        HeaderFontSize=None,
        Heading1FontSize=None,
        enFontName=None,
        zhFontName=None,
        titleSuffix=None,
        templatePath=None,
        extensions=None,
        directory=None,
    ) -> None:
        # 获取配置信息
        curPath = os.path.abspath(__file__)
        parentPath = os.path.dirname(os.path.dirname(curPath))
        fileName = "config.json"
        filePath = os.path.join(parentPath, fileName)
        with open(filePath, "r", encoding="utf-8") as configfile:
            try:
                config = json.load(configfile)
                infoConfig = config["info"]
                docxConfig = config["docx"]
                codeFileConfig = config["codeFile"]
                self.name = name if name is not (None or "") else infoConfig["name"]
                self.pageSize = pageSize if pageSize is not (None or 0) else docxConfig["pageSize"]
                self.HeaderFontSize = HeaderFontSize if HeaderFontSize is not (None or 0) else docxConfig["HeaderFontSize"]
                self.Heading1FontSize = Heading1FontSize if Heading1FontSize is not (None or 0) else docxConfig["Heading1FontSize"]
                self.enFontName = enFontName if enFontName is not (None or "") else docxConfig["enFontName"]
                self.zhFontName = zhFontName if zhFontName is not (None or "") else docxConfig["zhFontName"]
                self.titleSuffix = titleSuffix if titleSuffix is not (None or "") else docxConfig["titleSuffix"]
                self.templatePath = templatePath if templatePath is not (None or "") else docxConfig["templatePath"]
                self.extensions = extensions if extensions is not (None or [] or "") else codeFileConfig["extensions"]
                self.directory = directory if directory is not (None or "") else codeFileConfig["directory"]
                # self.pageSize = docxConfig["pageSize"]
                # self.HeaderFontSize = docxConfig["HeaderFontSize"]
                # self.Heading1FontSize = docxConfig["Heading1FontSize"]
                # self.enFontName = docxConfig["enFontName"]
                # self.zhFontName = docxConfig["zhFontName"]
                # self.titleSuffix = docxConfig["titleSuffix"]
                # self.templatePath = docxConfig["templatePath"]

                # self.extensions = codeFileConfig["extensions"]
                # self.directory = codeFileConfig["directory"]

                self.lineCount = 0
                fh = FileHandler(self.extensions, self.directory)
                self.files = fh.findFilesWithExtension(
                    extensions=self.extensions, directory=self.directory
                )
                for file in self.files:
                    with open(file, "r", encoding="utf-8") as f:
                        self.lineCount += len(f.readlines())
                self.needPartitions = self.lineCount > self.pageSize * 50
                print("总行数:", self.lineCount, ",是否需要分页:", self.needPartitions)

            except Exception as e:
                print("配置文件读取失败", e)
                exit()

        # 输出docx目标文件路径
        self.docxPath = os.path.join(
            parentPath,
            "out",
            str(self.name) + ".docx" if self.name is not None else "defult.docx",
        )

        # 初始化Document对象
        self.docxDocument = docx.Document(
            # "/Users/yushaochen/Code/Python/software-copyright-codehandler/res/template.docx"
            self.templatePath
        )
        # "/Users/yushaochen/Code/Python/software-copyright-codehandler/res/template.docx"

        # 设置字体，中文宋体，英文Times New Roman
        normalStyle = self.docxDocument.styles["Normal"]
        normalStyle.font.name = self.enFontName
        normalStyle.element.rPr.rFonts.set(qn("w:eastAsia"), self.zhFontName)

        headerStyle = self.docxDocument.styles["Header"]
        headerStyle.font.name = self.enFontName
        headerStyle.font.color.rgb = RGBColor(0, 0, 0)
        headerStyle.font.size = Pt(self.HeaderFontSize)
        headerStyle.element.rPr.rFonts.set(qn("w:eastAsia"), self.zhFontName)

        heading1Style = self.docxDocument.styles["Heading 1"]
        heading1Style.font.color.rgb = RGBColor(0, 0, 0)
        heading1Style.font.size = Pt(self.Heading1FontSize)

        # 设置标题
        self.addFirstPage()


    def process(self):
        """处理docx正文"""
        # 初始化文件处理对象，获取源代码文件列表
        fh = FileHandler()
        # 设置页眉
        # self.docxDocument.add_section(WD_SECTION_START.NEW_PAGE)
        header = self.docxDocument.sections[1].header
        header.is_linked_to_previous = False
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        headerRun = header.paragraphs[0].add_run(self.name + self.titleSuffix)
        headerRun.font.size = Pt(10.5)
        headerRun.font.name = self.enFontName
        headerRun.element.rPr.rFonts.set(qn("w:eastAsia"), self.zhFontName)

        # 通过wordxml对象设置页眉下边框
        bottom = {"sz": "12", "val": "single", "color": "#000000", "space": "0"}
        element = OxmlElement(f"w:bottom")
        element.set(qn("w:sz"), bottom["sz"])
        element.set(qn("w:val"), bottom["val"])
        element.set(qn("w:color"), bottom["color"])
        element.set(qn("w:space"), bottom["space"])
        he = header.paragraphs[0]._element
        hepr = he.get_or_add_pPr()
        hepr.append(element)

        # footer = self.docxDocument.sections[1].footer
        # footer.is_linked_to_previous = False
        # footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 生成docx正文
        if self.needPartitions:
            self.inputPartition()
        else:
            self.inputWithoutPartition()

        # 保存
        self.saveDocx()
        return self.docxPath

    def inputWithoutPartition(self):
        for i, file in enumerate(self.files):
            # 文件名
            heading1 = self.docxDocument.add_heading(level=1)
            run = heading1.add_run(os.path.basename(file) + ":")
            run.bold = True
            run.font.name = self.enFontName
            run.element.rPr.rFonts.set(qn("w:eastAsia"), self.zhFontName)
            heading1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # 源代码内容
            para = self.docxDocument.add_paragraph()
            para.style.font.size = Pt(9)

            with open(file, "r", encoding="utf-8") as f:
                for line in f.readlines():
                    para.add_run(line)

    def inputPartition(self):
        linesCount = 0
        flagOf30 = False
        for i, file in enumerate(self.files):
            if flagOf30:
                break
            # 文件名
            heading1 = self.docxDocument.add_heading(level=1)
            run = heading1.add_run(os.path.basename(file) + ":")
            run.bold = True
            run.font.name = self.enFontName
            run.element.rPr.rFonts.set(qn("w:eastAsia"), self.zhFontName)
            heading1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # 源代码内容
            para = self.docxDocument.add_paragraph()
            para.style.font.size = Pt(9)

            with open(file, "r", encoding="utf-8") as f:
                for line in f.readlines():
                    para.add_run(line)
                    linesCount += 1
                    if linesCount >= (self.pageSize / 2 - 1) * 50:
                        flagOf30 = True
                        break
        if flagOf30:
            self.files.reverse()
            flagOf30 = False
            linesCount = 0
            for i, file in enumerate(self.files):
                if flagOf30:
                    break

                # 源代码内容
                para = self.docxDocument.paragraphs[-1 - i].insert_paragraph_before()
                para.style.font.size = Pt(9)
                # 文件名
                filename = para.add_run(os.path.basename(file) + ":\n")
                filename.bold = True
                filename.font.name = self.enFontName
                filename.element.rPr.rFonts.set(qn("w:eastAsia"), self.zhFontName)
                filename.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                with open(file, "r", encoding="utf-8") as f:
                    # br = fh.BackwardsReader(f)
                    for line in f.readlines():
                        para.add_run(line)
                        linesCount += 1
                        if linesCount >= self.pageSize / 2 * 50:
                            flagOf30 = True
                            break

    def addFirstPage(self):
        """生成docx封面"""
        emptyLineNums = 5
        firstPageSection = self.docxDocument.sections[0]
        firstPageSection.header.is_linked_to_previous = False

        for i in range(emptyLineNums):
            header = self.docxDocument.paragraphs[0].insert_paragraph_before("")
            header.style = self.docxDocument.styles["Header"]
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        header = self.docxDocument.paragraphs[emptyLineNums].insert_paragraph_before(
            self.name + self.titleSuffix
        )
        header.style = self.docxDocument.styles["Header"]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # self.docxDocument.add_page_break()

    def saveDocx(self, docxPath=None):
        """保存docx文件到指定的路径

        Args:
            docxPath (str, optional): 指定的路径，不提供则使用配置文件中提供的路径. Defaults to None.
        """
        # 参数未提供时使用配置文件中的参数
        if docxPath is None:
            docxPath = self.docxPath
        # 初始化文件夹并保存docx
        if not os.path.exists(os.path.dirname(docxPath)):
            os.makedirs(os.path.dirname(docxPath))
        self.docxDocument.save(self.docxPath)


# dh = DocxHandler()
# dh.process()
